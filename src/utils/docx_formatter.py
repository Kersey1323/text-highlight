from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import re

def set_highlight_color(run, color):
    """
    Set highlight (background) color for a run.

    Args:
        run: docx run object
        color: RGBColor object
    """
    rPr = run._element.get_or_add_rPr()
    shading = OxmlElement('w:shd')
    # Convert RGBColor to hex string (Word expects RRGGBB)
    hex_color = str(color)
    shading.set(qn('w:fill'), hex_color)
    rPr.append(shading)

def add_text_with_highlights(paragraph, text, sensitive_words):
    """
    Add text to paragraph with highlighted sensitive words.

    Args:
        paragraph: docx paragraph object
        text: text to add
        sensitive_words: list of words to highlight
    """
    # Use regex to find all sensitive words and split the text
    if not sensitive_words:
        # No sensitive words, just add the whole text
        run = paragraph.add_run(text)
        _set_run_font(run)
        return

    # Sort words by length (desc) to handle overlaps/substrings
    sorted_words = sorted([w for w in sensitive_words if w.strip()], key=len, reverse=True)
    
    # Create a regex pattern to match any of the sensitive words
    # Escape special characters in words
    pattern = '|'.join(map(re.escape, sorted_words))
    
    last_idx = 0
    # Iterate through all matches
    for match in re.finditer(pattern, text):
        start, end = match.span()
        
        # Add text before match
        if start > last_idx:
            run = paragraph.add_run(text[last_idx:start])
            _set_run_font(run)
            
        # Add highlighted text
        run = paragraph.add_run(text[start:end])
        _set_run_font(run)
        set_highlight_color(run, RGBColor(255, 255, 0)) # Yellow
        
        last_idx = end
        
    # Add remaining text
    if last_idx < len(text):
        run = paragraph.add_run(text[last_idx:])
        _set_run_font(run)

def _set_run_font(run):
    # 字体设置：中文仿宋，英文Times New Roman
    run.font.name = '仿宋_GB2312'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run.font.size = Pt(12)  # 小四 = 12pt

def parse_markdown_to_docx(text, sensitive_words=None, doc=None):
    """
    Parse OCR markdown text and create a formatted Word document.
    Supports alignment tags: <center>...</center>, <right>...</right>

    Args:
        text: OCR text in markdown format
        sensitive_words: List of words to highlight (yellow background)
        doc: Optional existing Document object to append to. If None, a new one is created.

    Returns:
        Document: python-docx Document object
    """
    if doc is None:
        doc = Document()
        
        # Set margins to be slightly narrower to accommodate "Official Document" spacing
        # Standard Word margins are ~2.54cm or 3.17cm. We'll reduce to ~2.0cm.
        try:
            section = doc.sections[0]
            section.top_margin = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)
        except:
            pass

        # Set default font for normal style
        style = doc.styles['Normal']
        style.font.name = '仿宋_GB2312'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        style.font.size = Pt(12)

    if not sensitive_words:
        sensitive_words = []

    # Split text into lines and filter empty lines
    lines = text.split('\n')
    paragraphs = [line.strip() for line in lines if line.strip()]

    # Add paragraphs with single line spacing
    for line in paragraphs:
        # Check alignment tags
        alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # Default
        
        # Check for <center> or [CENTER]
        if re.search(r'<(center|CENTER)>', line) or re.search(r'\[(center|CENTER)\]', line):
            alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Clean tags
            line = re.sub(r'</?(center|CENTER)>', '', line)
            line = re.sub(r'\[/?(center|CENTER)\]', '', line)
            
        # Check for <right> or [RIGHT]
        elif re.search(r'<(right|RIGHT)>', line) or re.search(r'\[(right|RIGHT)\]', line):
            alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            # Clean tags
            line = re.sub(r'</?(right|RIGHT)>', '', line)
            line = re.sub(r'\[/?(right|RIGHT)\]', '', line)

        line = line.strip()
        if not line:
            continue

        # Check if it's a heading (markdown format: # Heading)
        if line.startswith('#'):
            # It's a heading
            heading_level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()

            # Add heading with appropriate style
            if heading_level == 1:
                # Main title - centered, larger font
                p = doc.add_paragraph(heading_text)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.runs[0]
                run.font.name = '仿宋_GB2312'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                run.font.size = Pt(22)
                run.bold = True
            elif heading_level == 2:
                # Secondary heading - 小标题格式
                p = doc.add_paragraph(heading_text)
                run = p.runs[0]
                # 字体：仿宋，英文Times New Roman
                run.font.name = '仿宋_GB2312'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                run.font.size = Pt(12)  # 小四 = 12pt
                run.font.bold = True  # 加粗
                # 首行缩进2字符
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.line_spacing = Pt(28)  # 固定值28磅，与正文一致
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
            else:
                # Other headings
                p = doc.add_paragraph(heading_text)
                run = p.runs[0]
                run.font.size = Pt(14)
                run.bold = True
        else:
            # It's a normal paragraph
            p = doc.add_paragraph()
            
            # Apply alignment detected earlier
            p.paragraph_format.alignment = alignment
            
            # Apply standard formatting if it's a normal justified paragraph
            if alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
                
            p.paragraph_format.line_spacing = Pt(28)  # 固定值28磅
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            add_text_with_highlights(p, line, sensitive_words)

    return doc

def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Convert Word document to PDF.

    Args:
        docx_path: Input Word document path
        pdf_path: Output PDF path
    """
    import time
    import pythoncom
    import os

    def _kill_word_processes():
        """Attempts to kill any lingering WINWORD.EXE processes to free COM locks."""
        try:
            import psutil
            for proc in psutil.process_iter(['pid', 'name']):
                if proc.info['name'] and proc.info['name'].lower() == 'winword.exe':
                    print(f"Killing lingering Word process: {proc.info['pid']}")
                    proc.kill()
        except ImportError:
            # Fallback if psutil is not installed
            import os
            os.system("taskkill /f /im winword.exe >nul 2>&1")
        except Exception:
            pass

    max_retries = 3
    for attempt in range(max_retries):
        try:
            # Initialize COM in this thread
            pythoncom.CoInitialize()
            
            # Ensure path is absolute for Word COM
            docx_path = os.path.abspath(docx_path)
            pdf_path = os.path.abspath(pdf_path)

            # Try using win32com directly for better control
            import win32com.client
            
            # Try to attach to existing or create new
            try:
                word = win32com.client.GetActiveObject("Word.Application")
            except Exception:
                word = win32com.client.Dispatch("Word.Application")
            
            word.Visible = False
            word.DisplayAlerts = False
            
            doc = word.Documents.Open(docx_path)
            try:
                doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF
                print(f"Saved PDF to {pdf_path}")
                return True
            finally:
                doc.Close(SaveChanges=0) # 0 = wdDoNotSaveChanges
                # Don't quit Word if we just attached to an existing instance that might be used elsewhere
                # But here we probably want to clean up if we launched it.
                # For safety in this script context, let's quit if we can to avoid lingering processes.
                try:
                    word.Quit()
                except:
                    pass
                    
        except Exception as e:
            print(f"Error converting to PDF (Attempt {attempt+1}/{max_retries}): {e}")
            # If it's a COM error, sometimes killing Word helps
            if attempt < max_retries - 1:
                print("Retrying...")
                _kill_word_processes()
                time.sleep(1)
            else:
                import traceback
                traceback.print_exc()
                return False
        finally:
             pythoncom.CoUninitialize()
             
    return False
