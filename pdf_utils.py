import fitz  # PyMuPDF
import os
import io
import time
import re
from PIL import Image
from ocr_utils import get_ocr_text
from text_renderer import render_text_with_highlights
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml import OxmlElement


def set_highlight_color(run, color):
    """
    Set highlight (background) color for a run.

    Args:
        run: docx run object
        color: RGBColor object
    """
    rPr = run._element.get_or_add_rPr()
    shading = OxmlElement('w:shd')
    # Convert RGBColor to hex string (Word expects RRGGBB)
    hex_color = "{:02X}{:02X}{:02X}".format(color.r, color.g, color.b)
    shading.set(qn('w:fill'), hex_color)
    rPr.append(shading)


def parse_markdown_to_docx(text, sensitive_words=None):
    """
    Parse OCR markdown text and create a formatted Word document.

    Args:
        text: OCR text in markdown format
        sensitive_words: List of words to highlight (yellow background)

    Returns:
        Document: python-docx Document object
    """
    doc = Document()

    # Set default font for normal style
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style.font.size = Pt(12)

    if not sensitive_words:
        sensitive_words = []

    # Split text into lines and filter empty lines
    lines = text.split('\n')
    paragraphs = [line.strip() for line in lines if line.strip()]

    # Add paragraphs with single line spacing
    for line in paragraphs:
        # Check if it's a heading (markdown format: # Heading)
        if line.startswith('#'):
            # It's a heading
            heading_level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()

            # Add heading with appropriate style
            if heading_level == 1:
                # Main title - centered, larger font
                p = doc.add_paragraph(heading_text)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.runs[0]
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(22)
                run.bold = True
            elif heading_level == 2:
                # Secondary heading
                p = doc.add_paragraph(heading_text)
                run = p.runs[0]
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(16)
                run.bold = True
            else:
                # Other headings
                p = doc.add_paragraph(heading_text)
                run = p.runs[0]
                run.font.size = Pt(14)
                run.bold = True
        else:
            # It's a normal paragraph
            # Check for sensitive words and highlight them
            p = doc.add_paragraph()
            add_text_with_highlights(p, line, sensitive_words)

    return doc


def add_text_with_highlights(paragraph, text, sensitive_words):
    """
    Add text to paragraph with highlighted sensitive words.

    Args:
        paragraph: docx paragraph object
        text: text to add
        sensitive_words: list of words to highlight
    """
    # For now, just add text without highlighting
    run = paragraph.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)


def ocr_and_highlight_pdf_docx(pdf_path, sensitive_words, output_docx_path, output_pdf_path=None):
    """
    Process PDF by OCR, create formatted Word document, and convert to PDF.

    Args:
        pdf_path: Input PDF path
        sensitive_words: List of words to highlight
        output_docx_path: Output Word document path
        output_pdf_path: Optional output PDF path (if None, only create docx)
    """
    try:
        doc = fitz.open(pdf_path)

        print(f"Processing PDF: {pdf_path} ({len(doc)} pages)")
        print("Mode: Create formatted Word document from OCR text")

        # Combine all OCR text
        all_text = ""

        for page_num, page in enumerate(doc):
            print(f"  Processing page {page_num + 1}/{len(doc)}...")

            # Convert to image for OCR
            pix = page.get_pixmap(dpi=200)
            project_root = os.path.dirname(os.path.abspath(__file__))
            temp_dir = os.path.join(project_root, "temp")
            os.makedirs(temp_dir, exist_ok=True)
            temp_img_path = os.path.join(temp_dir, f"temp_page_{page_num}.png")
            pix.save(temp_img_path)

            try:
                # OCR the image
                text = get_ocr_text(temp_img_path)

                if text:
                    print(f"    Text: {text}")
                    # Remove "识别结果：" prefix if present using regex
                    text = re.sub(r'^识别结果：', '', str(text)).strip()
                    all_text += text + "\n\n"
                else:
                    print(f"    Warning: No text found on page {page_num + 1}.")
            finally:
                # Cleanup temp file
                if os.path.exists(temp_img_path):
                    try:
                        os.remove(temp_img_path)
                    except:
                        pass

        # Create Word document
        print("Creating Word document...")
        docx_doc = parse_markdown_to_docx(all_text, sensitive_words)

        # Save Word document
        docx_doc.save(output_docx_path)
        print(f"Saved Word document to {output_docx_path}")

        # Convert to PDF if requested
        if output_pdf_path:
            print(f"Converting to PDF: {output_pdf_path}")
            convert_docx_to_pdf(output_docx_path, output_pdf_path)

        return True

    except Exception as e:
        print(f"Error processing PDF: {e}")
        import traceback
        traceback.print_exc()
        return False


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Convert Word document to PDF.

    Args:
        docx_path: Input Word document path
        pdf_path: Output PDF path
    """
    try:
        # Try using docx2pdf
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"Saved PDF to {pdf_path}")
    except ImportError:
        print("Warning: docx2pdf not installed. PDF conversion skipped.")
        print("Install with: pip install docx2pdf")
        # Alternative: try using win32com on Windows
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF
            doc.Close()
            word.Quit()
            print(f"Saved PDF to {pdf_path}")
        except ImportError:
            print("Warning: win32com not available. Please install docx2pdf for PDF conversion.")
        except Exception as e:
            print(f"Error converting to PDF: {e}")
    except Exception as e:
        print(f"Error converting to PDF: {e}")


def ocr_and_highlight_pdf_text_layer(pdf_path, sensitive_words, output_path):
    """
    Process PDF by OCR, then add text layer with highlights on top of original page.
    This creates a searchable PDF with text layer over the original scan.

    Args:
        pdf_path: Input PDF path
        sensitive_words: List of words to highlight
        output_path: Output PDF path
    """
    try:
        doc = fitz.open(pdf_path)
        output_doc = fitz.open()

        print(f"Processing PDF: {pdf_path} ({len(doc)} pages)")
        print("Mode: Adding text layer with highlights (preserves original layout)")

        # Text formatting configuration
        font_name = "china-s"  # Simplified Chinese font in PyMuPDF
        font_size = 12
        text_color = (0, 0, 0)  # Black
        highlight_color = (1, 1, 0)  # Yellow
        line_height_ratio = 1.5
        margin = 50  # Points from page edge

        for page_num, page in enumerate(doc):
            print(f"  Processing page {page_num + 1}/{len(doc)}...")

            # Get page dimensions
            rect = page.rect
            page_width = rect.width
            page_height = rect.height

            # 1. Convert to image for OCR
            pix = page.get_pixmap(dpi=200)
            project_root = os.path.dirname(os.path.abspath(__file__))
            temp_dir = os.path.join(project_root, "temp")
            os.makedirs(temp_dir, exist_ok=True)
            temp_img_path = os.path.join(temp_dir, f"temp_page_{page_num}.png")
            pix.save(temp_img_path)

            try:
                # 2. OCR the image
                text = get_ocr_text(temp_img_path)

                if not text:
                    print(f"    Warning: No text found on page {page_num + 1}.")
                    text = "[OCR Failed or Empty Page]"

                print(f"    OCR text length: {len(text)} chars")

                # 3. Create new page with original content as background
                new_page = output_doc.new_page(width=page_width, height=page_height)
                new_page.show_pdf_page(new_page.rect, doc, page_num)

                # 4. Add text layer with highlights
                # Parse text into paragraphs/lines
                paragraphs = parse_text_to_paragraphs(text)

                # Calculate starting Y position (from top)
                y_position = margin
                line_height = font_size * line_height_ratio

                for para in paragraphs:
                    # Check if paragraph fits on current page
                    if y_position + line_height > page_height - margin:
                        break  # Skip if doesn't fit

                    # Check for sensitive words in paragraph
                    highlighted_para, has_highlight = highlight_sensitive_words(
                        para, sensitive_words, new_page, margin, y_position,
                        font_name, font_size, text_color, highlight_color
                    )

                    if not has_highlight:
                        # No sensitive words, just add normal text
                        # Add semi-transparent white background for readability
                        text_rect = fitz.Rect(margin, y_position - 2,
                                            page_width - margin, y_position + line_height - 2)
                        new_page.draw_rect(text_rect, color=(0.8, 0.8, 0.8, 0.7), fill=(0.8, 0.8, 0.8, 0.7))

                        # Add text (truncate if too long)
                        display_text = para[:150] + "..." if len(para) > 150 else para
                        try:
                            new_page.insert_text(
                                (margin, y_position),
                                display_text,
                                fontname=font_name,
                                fontsize=font_size,
                                color=text_color
                            )
                        except Exception as e:
                            print(f"    Warning: Could not insert text: {e}")

                    y_position += line_height * 2  # Add some spacing between paragraphs

                print(f"    Added text layer to page {page_num + 1}")

            finally:
                # Cleanup temp file
                if os.path.exists(temp_img_path):
                    try:
                        os.remove(temp_img_path)
                    except:
                        pass

        # Save final PDF
        try:
            output_doc.save(output_path)
            print(f"Saved highlighted PDF to {output_path}")
        except Exception as e:
            if "Permission denied" in str(e) or "cannot remove file" in str(e):
                print(f"Warning: Could not save to {output_path} (File might be open).")
                base, ext = os.path.splitext(output_path)
                timestamp = int(time.time())
                new_path = f"{base}_{timestamp}{ext}"
                print(f"Attempting to save to {new_path} instead...")
                output_doc.save(new_path)
                print(f"Saved highlighted PDF to {new_path}")
            else:
                raise e

        return True

    except Exception as e:
        print(f"Error processing PDF: {e}")
        import traceback
        traceback.print_exc()
        return False


def parse_text_to_paragraphs(text):
    """Parse OCR text into paragraphs/lines."""
    # Split by newlines and filter empty lines
    lines = text.split('\n')
    paragraphs = [line.strip() for line in lines if line.strip()]
    return paragraphs


def highlight_sensitive_words(paragraph, sensitive_words, page, x, y,
                              font_name, font_size, text_color, highlight_color):
    """
    Add text with highlighted sensitive words to PDF page.

    Returns:
        tuple: (modified_paragraph, has_highlight)
    """
    has_highlight = False

    # Sort sensitive words by length (longest first) to avoid partial matches
    sorted_words = sorted([w for w in sensitive_words if w.strip()], key=len, reverse=True)

    # Check each sensitive word
    for word in sorted_words:
        if word in paragraph:
            has_highlight = True
            # Find all occurrences
            start = 0
            while True:
                idx = paragraph.find(word, start)
                if idx == -1:
                    break

                # Calculate approximate position (this is simplified)
                # In practice, you'd need more sophisticated positioning
                try:
                    # Add highlight annotation (approximate position)
                    # Note: This is a simplified approach
                    highlight_rect = fitz.Rect(x + idx * 6, y - 2,
                                             x + (idx + len(word)) * 6, y + font_size)
                    highlight = page.add_highlight_annot(highlight_rect)
                    highlight.set_colors(stroke=highlight_color)
                    highlight.update()
                except:
                    pass

                start = idx + len(word)

    return paragraph, has_highlight


def ocr_and_highlight_pdf(pdf_path, sensitive_words, output_path):
    """
    Process PDF by converting pages to images, running OCR,
    re-rendering text with highlights, and saving as a new PDF.
    """
    try:
        doc = fitz.open(pdf_path)
        output_doc = fitz.open()
        
        print(f"Processing PDF: {pdf_path} ({len(doc)} pages)")
        
        for page_num, page in enumerate(doc):
            print(f"  Processing page {page_num + 1}/{len(doc)}...")
            
            # 1. Convert Page to Image
            # Use higher DPI for better OCR accuracy
            pix = page.get_pixmap(dpi=300)

            # Save to a temporary file because our OCR utils expects a file path
            # Use a temp directory in the project folder
            project_root = os.path.dirname(os.path.abspath(__file__))
            temp_dir = os.path.join(project_root, "temp")
            os.makedirs(temp_dir, exist_ok=True)
            temp_img_path = os.path.join(temp_dir, f"temp_page_{page_num}.png")
            pix.save(temp_img_path)
            
            try:
                # 2. OCR the image
                text = get_ocr_text(temp_img_path)
                
                print(f"--- Page {page_num + 1} OCR Result Start ---")
                print(text)
                print(f"--- Page {page_num + 1} OCR Result End ---")
                
                if not text:
                    print(f"    Warning: No text found on page {page_num + 1}.")
                    text = "[OCR Failed or Empty Page]"
                
                # 3. Render text to new image with highlights
                # We don't save to file immediately, we get the PIL Image
                highlighted_img = render_text_with_highlights(text, sensitive_words, output_image_path=None)
                
                # 4. Convert PIL Image to bytes for PyMuPDF
                img_byte_arr = io.BytesIO()
                highlighted_img.save(img_byte_arr, format='PNG')
                img_bytes = img_byte_arr.getvalue()
                
                # 5. Create a new page in output PDF
                # We use the dimensions of the generated image
                output_page = output_doc.new_page(width=highlighted_img.width, height=highlighted_img.height)
                output_page.insert_image(output_page.rect, stream=img_bytes)
                
            finally:
                # Cleanup temp file
                if os.path.exists(temp_img_path):
                    try:
                        os.remove(temp_img_path)
                    except:
                        pass # Ignore cleanup errors
        
        # Save final PDF with retry logic for permission errors
        try:
            output_doc.save(output_path)
            print(f"Saved highlighted PDF to {output_path}")
        except Exception as e:
            if "Permission denied" in str(e) or "cannot remove file" in str(e):
                print(f"Warning: Could not save to {output_path} (File might be open).")
                # Try a new filename
                base, ext = os.path.splitext(output_path)
                timestamp = int(time.time())
                new_path = f"{base}_{timestamp}{ext}"
                print(f"Attempting to save to {new_path} instead...")
                output_doc.save(new_path)
                print(f"Saved highlighted PDF to {new_path}")
            else:
                raise e
                
        return True

    except Exception as e:
        print(f"Error processing PDF: {e}")
        import traceback
        traceback.print_exc()
        return False
